#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MultiConvert Pro - Conversor de Documentos

Este módulo contém a classe DocumentConverter que gerencia
conversões de arquivos de documentos usando LibreOffice CLI
e bibliotecas Python como fallback.

Autor: MultiConvert Pro Team
Versão: 1.0.0
"""

import os
import subprocess
from typing import Optional, Callable
from pathlib import Path

from ..engines.onlyoffice_engine import OnlyOfficeEngine
from ..engines.libreoffice_engine import LibreOfficeEngine
from ..engines.fallback_engine import FallbackEngine


class DocumentConverter:
    """Conversor especializado para arquivos de documentos com sistema de fallback."""
    
    def __init__(self):
        self.supported_formats = {
            'input': ['pdf', 'docx', 'doc', 'odt', 'rtf', 'txt', 'xlsx', 'xls', 'ods', 'pptx', 'ppt', 'odp'],
            'output': ['pdf', 'docx', 'odt', 'rtf', 'txt', 'xlsx', 'ods', 'pptx', 'odp', 'html']
        }
        
        # Inicializar engines
        self.onlyoffice_engine = OnlyOfficeEngine()
        self.libreoffice_engine = LibreOfficeEngine()
        self.fallback_engine = FallbackEngine()
        
        # Mapeamento de conversões suportadas por cada engine
        self.engine_capabilities = {
            'onlyoffice': {
                'input': ['txt', 'docx', 'doc', 'odt', 'rtf', 'xlsx', 'xls', 'ods', 'pptx', 'ppt', 'odp'],
                'output': ['pdf', 'docx', 'odt', 'rtf', 'txt', 'html', 'xlsx', 'ods', 'pptx', 'odp']
            },
            'libreoffice': {
                'input': ['docx', 'doc', 'odt', 'rtf', 'txt', 'xlsx', 'xls', 'ods', 'pptx', 'ppt', 'odp'],
                'output': ['pdf', 'docx', 'odt', 'rtf', 'txt', 'xlsx', 'ods', 'pptx', 'odp', 'html']
            },
            'fallback': {
                'input': ['pdf', 'docx', 'txt', 'rtf'],
                'output': ['txt', 'docx', 'pdf', 'html']
            }
        }
        
        # Mapeamento de conversões para uma lista de engines (ordem de prioridade)
        # Nota: OnlyOffice será usado apenas se estiver disponível
        available_engines = []
        if self.onlyoffice_engine.is_available():
            available_engines.append(self.onlyoffice_engine)
        if self.libreoffice_engine.is_available():
            available_engines.append(self.libreoffice_engine)
        
        self.conversion_methods = {
            ('txt', 'pdf'): available_engines.copy(),
            ('docx', 'pdf'): available_engines.copy(),
            ('doc', 'pdf'): available_engines.copy(),
            ('odt', 'pdf'): available_engines.copy(),
            ('rtf', 'pdf'): available_engines.copy(),
            ('xlsx', 'pdf'): available_engines.copy(),
            ('xls', 'pdf'): available_engines.copy(),
            ('ods', 'pdf'): available_engines.copy(),
            ('pptx', 'pdf'): available_engines.copy(),
            ('ppt', 'pdf'): available_engines.copy(),
            ('odp', 'pdf'): available_engines.copy(),
            ('pdf', 'docx'): [self.fallback_engine],
            ('pdf', 'txt'): [self.fallback_engine],
            ('docx', 'txt'): available_engines + [self.fallback_engine],
            ('txt', 'docx'): available_engines + [self.fallback_engine]
        }
        
        # Presets de qualidade para documentos
        self.quality_presets = {
            'baixa': {
                'pdf_quality': 'screen',
                'image_dpi': 72,
                'compress_images': True
            },
            'media': {
                'pdf_quality': 'print',
                'image_dpi': 150,
                'compress_images': True
            },
            'alta': {
                'pdf_quality': 'prepress',
                'image_dpi': 300,
                'compress_images': False
            },
            'maxima': {
                'pdf_quality': 'prepress',
                'image_dpi': 600,
                'compress_images': False
            }
        }
    
    def is_available(self) -> bool:
        """Verifica se o conversor de documentos está disponível."""
        # Verificar se pelo menos uma das engines está disponível
        try:
            from ..engines.onlyoffice_engine import OnlyOfficeEngine
            from ..engines.libreoffice_engine import LibreOfficeEngine
            from ..engines.fallback_engine import FallbackEngine
            
            onlyoffice = OnlyOfficeEngine()
            libreoffice = LibreOfficeEngine()
            fallback = FallbackEngine()
            
            return onlyoffice.is_available() or libreoffice.is_available() or fallback.is_available()
        except Exception:
            return False
    
    def is_supported_input(self, file_path: str) -> bool:
        """Verifica se o formato de entrada é suportado."""
        extension = Path(file_path).suffix.lower().lstrip('.')
        return extension in self.supported_formats['input']
    
    def is_supported_output(self, format_name: str) -> bool:
        """Verifica se o formato de saída é suportado."""
        return format_name.lower() in self.supported_formats['output']
    
    def can_convert(self, input_format: str, output_format: str) -> bool:
        """Verifica se uma conversão específica é suportada."""
        return (input_format.lower() in self.supported_formats['input'] and 
                output_format.lower() in self.supported_formats['output'])
    
    def get_supported_input_formats(self) -> list:
        """Retorna lista de formatos de entrada suportados."""
        return self.supported_formats['input']
    
    def get_supported_output_formats(self) -> list:
        """Retorna lista de formatos de saída suportados."""
        return self.supported_formats['output']
    
    def get_engines_status(self) -> dict:
        """Retorna o status de todos os engines disponíveis."""
        return {
            'onlyoffice': {
                'available': self.onlyoffice_engine.is_available(),
                'version': self.onlyoffice_engine.get_version(),
                'description': 'Conversão de alta qualidade com OnlyOffice DocumentBuilder'
            },
            'libreoffice': {
                'available': self.libreoffice_engine.is_available(),
                'version': self.libreoffice_engine.get_version(),
                'description': 'Conversão de alta fidelidade'
            },
            'fallback': {
                'available': self.fallback_engine.is_available(),
                'version': 'Bibliotecas Python',
                'description': 'Conversão básica com bibliotecas Python'
            }
        }
    
    def get_file_info(self, file_path: str) -> dict:
        """Obtém informações básicas do arquivo de documento."""
        try:
            stat = os.stat(file_path)
            extension = Path(file_path).suffix.lower().lstrip('.')
            
            info = {
                'size_bytes': stat.st_size,
                'extension': extension,
                'type': self._get_document_type(extension),
                'created': stat.st_ctime,
                'modified': stat.st_mtime
            }
            
            # Tentar obter informações específicas do formato
            if extension == 'pdf':
                info.update(self._get_pdf_info(file_path))
            elif extension in ['docx', 'doc']:
                info.update(self._get_word_info(file_path))
            
            return info
            
        except Exception as e:
            return {
                'error': f'Erro ao obter informações: {str(e)}',
                'size_bytes': 0,
                'extension': 'unknown',
                'type': 'unknown'
            }
    
    def _get_document_type(self, extension: str) -> str:
        """Determina o tipo de documento baseado na extensão."""
        type_mapping = {
            'pdf': 'document',
            'docx': 'word_document',
            'doc': 'word_document',
            'odt': 'writer_document',
            'rtf': 'rich_text',
            'txt': 'plain_text',
            'xlsx': 'spreadsheet',
            'xls': 'spreadsheet',
            'ods': 'calc_spreadsheet',
            'pptx': 'presentation',
            'ppt': 'presentation',
            'odp': 'impress_presentation'
        }
        return type_mapping.get(extension, 'unknown')
    
    def _get_pdf_info(self, file_path: str) -> dict:
        """Obtém informações específicas de arquivos PDF."""
        try:
            # Tentar usar PyPDF2 para informações básicas
            import PyPDF2
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                return {
                    'pages': len(reader.pages),
                    'encrypted': reader.is_encrypted,
                    'metadata': dict(reader.metadata) if reader.metadata else {}
                }
        except Exception:
            return {'pages': 'unknown', 'encrypted': False}
    
    def _get_word_info(self, file_path: str) -> dict:
        """Obtém informações específicas de arquivos Word."""
        try:
            # Tentar usar python-docx para informações básicas
            from docx import Document
            doc = Document(file_path)
            return {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'images': len([r for r in doc.part.rels.values() if 'image' in r.target_ref])
            }
        except Exception:
            return {'paragraphs': 'unknown', 'tables': 'unknown', 'images': 'unknown'}
    
    def convert(
        self,
        input_path: str,
        output_path: str,
        target_format: str,
        quality: str = 'media',
        progress_callback: Optional[Callable] = None
    ) -> tuple[bool, str]:
        """Converte um arquivo de documento usando sistema de fallback.
        
        Args:
            input_path: Caminho do arquivo de entrada
            output_path: Caminho do arquivo de saída
            target_format: Formato de saída (pdf, docx, etc.)
            quality: Preset de qualidade (baixa, media, alta, maxima)
            progress_callback: Callback para progresso
            
        Returns:
            Tupla (sucesso, mensagem)
        """
        try:
            # Validar entrada
            if not self.is_supported_input(input_path):
                return False, f"Formato de entrada não suportado: {Path(input_path).suffix}"
            
            if not self.is_supported_output(target_format):
                return False, f"Formato de saída não suportado: {target_format}"
            
            if not os.path.exists(input_path):
                return False, f"Arquivo não encontrado: {input_path}"
            
            input_ext = Path(input_path).suffix.lower().lstrip('.')
            
            if progress_callback:
                progress_callback(10, "Iniciando conversão de documento...")
            
            # Usar nova lógica de seleção de motor baseada em prioridade
            conversion_key = (input_ext, target_format)
            engines_to_try = self.conversion_methods.get(conversion_key, [])
            
            if not engines_to_try:
                # Fallback para lógica antiga se não houver mapeamento específico
                if self._can_use_onlyoffice(input_ext, target_format):
                    engines_to_try = [self.onlyoffice_engine, self.libreoffice_engine]
                elif self._can_use_libreoffice(input_ext, target_format):
                    engines_to_try = [self.libreoffice_engine]
                elif self._can_use_fallback(input_ext, target_format):
                    engines_to_try = [self.fallback_engine]
            
            if not engines_to_try:
                return False, f"Conversão de {input_ext} para {target_format} não suportada pelos engines disponíveis"
            
            last_error = ""
            progress_step = 80 // len(engines_to_try)  # Dividir progresso entre engines
            current_progress = 20
            
            # Tentar cada engine na ordem de prioridade
            for i, engine in enumerate(engines_to_try):
                engine_name = engine.__class__.__name__.replace('Engine', '')
                
                if progress_callback:
                    progress_callback(current_progress, f"Tentando conversão com {engine_name}...")
                
                try:
                    if engine == self.fallback_engine:
                        # FallbackEngine usa parâmetros diferentes
                        success, message = engine.convert(
                            input_path=input_path,
                            output_path=output_path,
                            input_format=input_ext,
                            output_format=target_format,
                            progress_callback=lambda p, m: progress_callback(current_progress + p * progress_step / 100, m) if progress_callback else None
                        )
                    else:
                        # OnlyOffice e LibreOffice usam parâmetros padrão
                        success, message = engine.convert(
                            input_path=input_path,
                            output_path=output_path,
                            target_format=target_format,
                            quality=quality,
                            progress_callback=lambda p, m: progress_callback(current_progress + p * progress_step / 100, m) if progress_callback else None
                        )
                    
                    if success:
                        if progress_callback:
                            progress_callback(100, f"Conversão concluída com {engine_name}!")
                        return True, f"Documento convertido com sucesso usando {engine_name}: {message}"
                    else:
                        last_error = message
                        if progress_callback and i < len(engines_to_try) - 1:
                            progress_callback(current_progress + progress_step, f"{engine_name} falhou, tentando próximo método...")
                
                except Exception as e:
                    last_error = f"Erro no {engine_name}: {str(e)}"
                    if progress_callback and i < len(engines_to_try) - 1:
                        progress_callback(current_progress + progress_step, f"Erro no {engine_name}, tentando próximo método...")
                
                current_progress += progress_step
            
            return False, f"Falha em todos os métodos de conversão. Último erro: {last_error}"
            
        except Exception as e:
            error_msg = f"Erro na conversão de documento: {str(e)}"
            return False, error_msg
    
    def _can_use_onlyoffice(self, input_ext: str, target_format: str) -> bool:
        """Verifica se o OnlyOffice pode fazer a conversão."""
        capabilities = self.engine_capabilities['onlyoffice']
        return (input_ext in capabilities['input'] and 
                target_format in capabilities['output'])
    
    def _can_use_libreoffice(self, input_ext: str, target_format: str) -> bool:
        """Verifica se o LibreOffice pode fazer a conversão."""
        capabilities = self.engine_capabilities['libreoffice']
        return (input_ext in capabilities['input'] and 
                target_format in capabilities['output'])
    
    def _can_use_fallback(self, input_ext: str, target_format: str) -> bool:
        """Verifica se o engine de fallback pode fazer a conversão."""
        capabilities = self.engine_capabilities['fallback']
        return (input_ext in capabilities['input'] and 
                target_format in capabilities['output'])
    
    def extract_text(self, input_path: str, output_path: str) -> tuple[bool, str]:
        """Extrai texto de um documento para arquivo TXT."""
        try:
            success, message = self.fallback_engine.extract_text(input_path, output_path)
            return success, message
        except Exception as e:
            return False, f"Erro na extração de texto: {str(e)}"
    
    def merge_pdfs(self, input_paths: list, output_path: str) -> tuple[bool, str]:
        """Mescla múltiplos PDFs em um único arquivo."""
        try:
            success, message = self.fallback_engine.merge_pdfs(input_paths, output_path)
            return success, message
        except Exception as e:
            return False, f"Erro na mesclagem de PDFs: {str(e)}"
    
    def split_pdf(self, input_path: str, output_dir: str, pages_per_file: int = 1) -> tuple[bool, str]:
        """Divide um PDF em múltiplos arquivos."""
        try:
            success, message = self.fallback_engine.split_pdf(input_path, output_dir, pages_per_file)
            return success, message
        except Exception as e:
            return False, f"Erro na divisão de PDF: {str(e)}"
    
    def get_recommended_settings(self, input_path: str, target_format: str) -> dict:
        """Retorna configurações recomendadas baseadas no arquivo de entrada."""
        try:
            info = self.get_file_info(input_path)
            input_ext = Path(input_path).suffix.lower().lstrip('.')
            
            # Configurações padrão
            settings = {
                'quality': 'media',
                'preserve_formatting': True,
                'preserve_metadata': True
            }
            
            # Ajustar qualidade baseada no tamanho do arquivo
            size_mb = info.get('size_bytes', 0) / (1024 * 1024)
            if size_mb > 50:  # Arquivo grande
                settings['quality'] = 'baixa'  # Comprimir mais
            elif size_mb < 1:  # Arquivo pequeno
                settings['quality'] = 'alta'  # Manter qualidade
            
            # Recomendações específicas por conversão
            conversion_recommendations = {
                ('pdf', 'docx'): {
                    'method': 'fallback',
                    'warning': 'Formatação pode ser alterada',
                    'preserve_images': True
                },
                ('docx', 'pdf'): {
                    'method': 'libreoffice',
                    'preserve_formatting': True,
                    'embed_fonts': True
                },
                ('txt', 'pdf'): {
                    'method': 'libreoffice',
                    'font': 'Arial',
                    'font_size': 12
                }
            }
            
            key = (input_ext, target_format)
            if key in conversion_recommendations:
                settings.update(conversion_recommendations[key])
            
            # Verificar disponibilidade dos engines
            if not self.libreoffice_engine.is_available():
                settings['libreoffice_warning'] = 'LibreOffice não encontrado, usando método alternativo'
            
            return settings
            
        except Exception:
            return {
                'quality': 'media',
                'preserve_formatting': True,
                'preserve_metadata': True
            }
    
    def get_conversion_strategy(self, input_ext: str, target_format: str) -> list:
        """Retorna a estratégia de conversão (ordem dos engines a tentar)."""
        strategies = []
        
        # Verificar se há mapeamento específico para esta conversão
        conversion_key = (input_ext, target_format)
        if conversion_key in self.conversion_methods:
            engines_list = self.conversion_methods[conversion_key]
            for engine in engines_list:
                engine_name = engine.__class__.__name__.replace('Engine', '').lower()
                if engine_name == 'onlyoffice':
                    strategies.append({
                        'engine': 'onlyoffice',
                        'quality': 'high',
                        'description': 'Conversão de alta qualidade com OnlyOffice DocumentBuilder'
                    })
                elif engine_name == 'libreoffice':
                    strategies.append({
                        'engine': 'libreoffice',
                        'quality': 'high',
                        'description': 'Conversão de alta fidelidade com LibreOffice'
                    })
                elif engine_name == 'fallback':
                    strategies.append({
                        'engine': 'fallback',
                        'quality': 'medium',
                        'description': 'Conversão básica com bibliotecas Python'
                    })
        else:
            # Fallback para lógica antiga
            # Adicionar OnlyOffice se suportado
            if self._can_use_onlyoffice(input_ext, target_format):
                strategies.append({
                    'engine': 'onlyoffice',
                    'quality': 'high',
                    'description': 'Conversão de alta qualidade com OnlyOffice DocumentBuilder'
                })
            
            # Adicionar LibreOffice se suportado
            if self._can_use_libreoffice(input_ext, target_format):
                strategies.append({
                    'engine': 'libreoffice',
                    'quality': 'high',
                    'description': 'Conversão de alta fidelidade com LibreOffice'
                })
            
            # Adicionar fallback se suportado
            if self._can_use_fallback(input_ext, target_format):
                strategies.append({
                    'engine': 'fallback',
                    'quality': 'medium',
                    'description': 'Conversão básica com bibliotecas Python'
                })
        
        return strategies