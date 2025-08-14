#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MultiConvert Pro - Engine de Fallback

Este módulo contém a classe FallbackEngine que executa
conversões de documentos usando bibliotecas Python como
fallback quando o LibreOffice não está disponível.

Autor: MultiConvert Pro Team
Versão: 1.0.0
"""

import os
import io
import tempfile
from typing import Optional, Callable, Dict, Any
from pathlib import Path

# Imports condicionais para bibliotecas de fallback
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    from pdf2docx import Converter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from striprtf.striprtf import rtf_to_text
    STRIPRTF_AVAILABLE = True
except ImportError:
    STRIPRTF_AVAILABLE = False


class FallbackEngine:
    """Engine de fallback para conversões usando bibliotecas Python."""
    
    def __init__(self):
        self.available_libraries = {
            'PyPDF2': PYPDF2_AVAILABLE,
            'python-docx': PYTHON_DOCX_AVAILABLE,
            'pdf2docx': PDF2DOCX_AVAILABLE,
            'reportlab': REPORTLAB_AVAILABLE,
            'striprtf': STRIPRTF_AVAILABLE
        }
        
        # Mapeamento de conversões suportadas
        self.conversion_matrix = {
            ('pdf', 'txt'): self._pdf_to_text,
            ('pdf', 'docx'): self._pdf_to_docx,
            ('docx', 'txt'): self._docx_to_text,
            ('docx', 'pdf'): self._docx_to_pdf,
            ('txt', 'pdf'): self._text_to_pdf,
            ('txt', 'docx'): self._text_to_docx,
            ('rtf', 'txt'): self._rtf_to_text,
            ('rtf', 'docx'): self._rtf_to_docx,
        }
    
    def is_available(self) -> bool:
        """Verifica se pelo menos uma biblioteca está disponível."""
        return any(self.available_libraries.values())
    
    def get_available_libraries(self) -> Dict[str, bool]:
        """Retorna o status das bibliotecas disponíveis."""
        return self.available_libraries.copy()
    
    def can_convert(self, input_format: str, output_format: str) -> bool:
        """Verifica se uma conversão específica é suportada."""
        conversion_key = (input_format.lower(), output_format.lower())
        return conversion_key in self.conversion_matrix
    
    def convert(
        self,
        input_path: str,
        output_path: str,
        input_format: str,
        output_format: str,
        progress_callback: Optional[Callable] = None
    ) -> tuple[bool, str]:
        """Converte um documento usando bibliotecas Python.
        
        Args:
            input_path: Caminho do arquivo de entrada
            output_path: Caminho do arquivo de saída
            input_format: Formato de entrada
            output_format: Formato de saída
            progress_callback: Callback para progresso
            
        Returns:
            Tupla (sucesso, mensagem)
        """
        try:
            if progress_callback:
                progress_callback(10, "Iniciando conversão com bibliotecas Python...")
            
            # Verificar se a conversão é suportada
            conversion_key = (input_format.lower(), output_format.lower())
            if conversion_key not in self.conversion_matrix:
                return False, f"Conversão {input_format} → {output_format} não suportada pelo engine de fallback"
            
            # Verificar se o arquivo de entrada existe
            if not os.path.exists(input_path):
                return False, f"Arquivo de entrada não encontrado: {input_path}"
            
            if progress_callback:
                progress_callback(20, "Preparando conversão...")
            
            # Criar diretório de saída se necessário
            output_dir = os.path.dirname(output_path)
            if output_dir:  # Só cria se não for string vazia (diretório atual)
                os.makedirs(output_dir, exist_ok=True)
            
            # Executar conversão específica
            converter_func = self.conversion_matrix[conversion_key]
            success, message = converter_func(input_path, output_path, progress_callback)
            
            if success and progress_callback:
                progress_callback(100, "Conversão concluída!")
            
            return success, message
            
        except Exception as e:
            return False, f"Erro na conversão com engine de fallback: {str(e)}"
    
    def _pdf_to_text(self, input_path: str, output_path: str, progress_callback: Optional[Callable] = None) -> tuple[bool, str]:
        """Converte PDF para texto usando PyPDF2."""
        if not PYPDF2_AVAILABLE:
            return False, "PyPDF2 não está instalado"
        
        try:
            if progress_callback:
                progress_callback(30, "Extraindo texto do PDF...")
            
            text_content = []
            
            with open(input_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                for i, page in enumerate(pdf_reader.pages):
                    if progress_callback:
                        page_progress = 30 + int((i / total_pages) * 50)
                        progress_callback(page_progress, f"Processando página {i+1} de {total_pages}...")
                    
                    text_content.append(page.extract_text())
            
            if progress_callback:
                progress_callback(85, "Salvando arquivo de texto...")
            
            # Salvar texto extraído
            with open(output_path, 'w', encoding='utf-8') as output_file:
                output_file.write('\n\n'.join(text_content))
            
            return True, "PDF convertido para texto com sucesso"
            
        except Exception as e:
            return False, f"Erro ao converter PDF para texto: {str(e)}"
    
    def _pdf_to_docx(self, input_path: str, output_path: str, progress_callback: Optional[Callable] = None) -> tuple[bool, str]:
        """Converte PDF para DOCX usando pdf2docx."""
        if not PDF2DOCX_AVAILABLE:
            return False, "pdf2docx não está instalado"
        
        try:
            if progress_callback:
                progress_callback(30, "Convertendo PDF para DOCX...")
            
            # Usar pdf2docx para conversão
            cv = Converter(input_path)
            cv.convert(output_path, start=0, end=None)
            cv.close()
            
            if progress_callback:
                progress_callback(90, "Finalizando conversão...")
            
            return True, "PDF convertido para DOCX com sucesso"
            
        except Exception as e:
            return False, f"Erro ao converter PDF para DOCX: {str(e)}"
    
    def _docx_to_text(self, input_path: str, output_path: str, progress_callback: Optional[Callable] = None) -> tuple[bool, str]:
        """Converte DOCX para texto usando python-docx."""
        if not PYTHON_DOCX_AVAILABLE:
            return False, "python-docx não está instalado"
        
        try:
            if progress_callback:
                progress_callback(30, "Extraindo texto do DOCX...")
            
            # Abrir documento DOCX
            doc = Document(input_path)
            text_content = []
            
            # Extrair texto dos parágrafos
            total_paragraphs = len(doc.paragraphs)
            for i, paragraph in enumerate(doc.paragraphs):
                if progress_callback and total_paragraphs > 0:
                    para_progress = 30 + int((i / total_paragraphs) * 50)
                    progress_callback(para_progress, f"Processando parágrafo {i+1} de {total_paragraphs}...")
                
                text_content.append(paragraph.text)
            
            # Extrair texto das tabelas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content.append(cell.text)
            
            if progress_callback:
                progress_callback(85, "Salvando arquivo de texto...")
            
            # Salvar texto extraído
            with open(output_path, 'w', encoding='utf-8') as output_file:
                output_file.write('\n'.join(text_content))
            
            return True, "DOCX convertido para texto com sucesso"
            
        except Exception as e:
            return False, f"Erro ao converter DOCX para texto: {str(e)}"
    
    def _docx_to_pdf(self, input_path: str, output_path: str, progress_callback: Optional[Callable] = None) -> tuple[bool, str]:
        """Converte DOCX para PDF usando python-docx + reportlab."""
        if not (PYTHON_DOCX_AVAILABLE and REPORTLAB_AVAILABLE):
            return False, "python-docx e/ou reportlab não estão instalados"
        
        try:
            if progress_callback:
                progress_callback(30, "Extraindo conteúdo do DOCX...")
            
            # Extrair texto do DOCX
            doc = Document(input_path)
            
            if progress_callback:
                progress_callback(50, "Criando PDF...")
            
            # Criar PDF com reportlab
            pdf_doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Adicionar parágrafos ao PDF
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    para = Paragraph(paragraph.text, styles['Normal'])
                    story.append(para)
                    story.append(Spacer(1, 12))
            
            if progress_callback:
                progress_callback(80, "Finalizando PDF...")
            
            # Construir PDF
            pdf_doc.build(story)
            
            return True, "DOCX convertido para PDF com sucesso (formatação básica)"
            
        except Exception as e:
            return False, f"Erro ao converter DOCX para PDF: {str(e)}"
    
    def _text_to_pdf(self, input_path: str, output_path: str, progress_callback: Optional[Callable] = None) -> tuple[bool, str]:
        """Converte texto para PDF usando reportlab."""
        if not REPORTLAB_AVAILABLE:
            return False, "reportlab não está instalado"
        
        try:
            if progress_callback:
                progress_callback(30, "Lendo arquivo de texto...")
            
            # Ler arquivo de texto
            with open(input_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            if progress_callback:
                progress_callback(50, "Criando PDF...")
            
            # Criar PDF
            pdf_doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Dividir texto em parágrafos
            paragraphs = text_content.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = Paragraph(para_text, styles['Normal'])
                    story.append(para)
                    story.append(Spacer(1, 12))
            
            if progress_callback:
                progress_callback(80, "Finalizando PDF...")
            
            # Construir PDF
            pdf_doc.build(story)
            
            return True, "Texto convertido para PDF com sucesso"
            
        except Exception as e:
            return False, f"Erro ao converter texto para PDF: {str(e)}"
    
    def _text_to_docx(self, input_path: str, output_path: str, progress_callback: Optional[Callable] = None) -> tuple[bool, str]:
        """Converte texto para DOCX usando python-docx."""
        if not PYTHON_DOCX_AVAILABLE:
            return False, "python-docx não está instalado"
        
        try:
            if progress_callback:
                progress_callback(30, "Lendo arquivo de texto...")
            
            # Ler arquivo de texto
            with open(input_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            if progress_callback:
                progress_callback(50, "Criando documento DOCX...")
            
            # Criar documento DOCX
            doc = Document()
            
            # Adicionar parágrafos
            paragraphs = text_content.split('\n')
            for para_text in paragraphs:
                doc.add_paragraph(para_text)
            
            if progress_callback:
                progress_callback(80, "Salvando documento...")
            
            # Salvar documento
            doc.save(output_path)
            
            return True, "Texto convertido para DOCX com sucesso"
            
        except Exception as e:
            return False, f"Erro ao converter texto para DOCX: {str(e)}"
    
    def _rtf_to_text(self, input_path: str, output_path: str, progress_callback: Optional[Callable] = None) -> tuple[bool, str]:
        """Converte RTF para texto usando striprtf."""
        if not STRIPRTF_AVAILABLE:
            return False, "striprtf não está instalado"
        
        try:
            if progress_callback:
                progress_callback(30, "Lendo arquivo RTF...")
            
            # Ler arquivo RTF
            with open(input_path, 'r', encoding='utf-8') as file:
                rtf_content = file.read()
            
            if progress_callback:
                progress_callback(60, "Convertendo RTF para texto...")
            
            # Converter RTF para texto
            text_content = rtf_to_text(rtf_content)
            
            if progress_callback:
                progress_callback(80, "Salvando arquivo de texto...")
            
            # Salvar texto
            with open(output_path, 'w', encoding='utf-8') as output_file:
                output_file.write(text_content)
            
            return True, "RTF convertido para texto com sucesso"
            
        except Exception as e:
            return False, f"Erro ao converter RTF para texto: {str(e)}"
    
    def _rtf_to_docx(self, input_path: str, output_path: str, progress_callback: Optional[Callable] = None) -> tuple[bool, str]:
        """Converte RTF para DOCX via texto intermediário."""
        if not (STRIPRTF_AVAILABLE and PYTHON_DOCX_AVAILABLE):
            return False, "striprtf e/ou python-docx não estão instalados"
        
        try:
            # Primeiro converter RTF para texto
            with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as temp_file:
                temp_path = temp_file.name
            
            success, message = self._rtf_to_text(input_path, temp_path, progress_callback)
            if not success:
                return False, message
            
            # Depois converter texto para DOCX
            success, message = self._text_to_docx(temp_path, output_path, progress_callback)
            
            # Limpar arquivo temporário
            try:
                os.unlink(temp_path)
            except Exception:
                pass
            
            return success, message
            
        except Exception as e:
            return False, f"Erro ao converter RTF para DOCX: {str(e)}"
    
    def get_supported_conversions(self) -> list:
        """Retorna lista de conversões suportadas."""
        supported = []
        for (input_fmt, output_fmt), func in self.conversion_matrix.items():
            # Verificar se as bibliotecas necessárias estão disponíveis
            try:
                # Teste rápido para ver se a função pode ser executada
                if input_fmt == 'pdf' and not PYPDF2_AVAILABLE and not PDF2DOCX_AVAILABLE:
                    continue
                if input_fmt == 'docx' and not PYTHON_DOCX_AVAILABLE:
                    continue
                if output_fmt == 'pdf' and not REPORTLAB_AVAILABLE:
                    continue
                if output_fmt == 'docx' and not PYTHON_DOCX_AVAILABLE:
                    continue
                if input_fmt == 'rtf' and not STRIPRTF_AVAILABLE:
                    continue
                
                supported.append({
                    'input': input_fmt,
                    'output': output_fmt,
                    'description': f"{input_fmt.upper()} → {output_fmt.upper()}"
                })
            except Exception:
                continue
        
        return supported
    
    def test_libraries(self) -> Dict[str, Any]:
        """Testa a disponibilidade e funcionalidade das bibliotecas."""
        results = {
            'available_libraries': self.available_libraries.copy(),
            'supported_conversions': len(self.get_supported_conversions()),
            'test_results': {}
        }
        
        # Teste básico para cada biblioteca
        if PYPDF2_AVAILABLE:
            try:
                # Teste simples do PyPDF2
                reader = PyPDF2.PdfReader
                results['test_results']['PyPDF2'] = 'OK'
            except Exception as e:
                results['test_results']['PyPDF2'] = f'Erro: {str(e)}'
        
        if PYTHON_DOCX_AVAILABLE:
            try:
                # Teste simples do python-docx
                doc = Document()
                results['test_results']['python-docx'] = 'OK'
            except Exception as e:
                results['test_results']['python-docx'] = f'Erro: {str(e)}'
        
        if REPORTLAB_AVAILABLE:
            try:
                # Teste simples do reportlab
                styles = getSampleStyleSheet()
                results['test_results']['reportlab'] = 'OK'
            except Exception as e:
                results['test_results']['reportlab'] = f'Erro: {str(e)}'
        
        return results